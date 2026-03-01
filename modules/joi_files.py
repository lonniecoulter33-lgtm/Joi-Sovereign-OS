"""
modules/joi_files.py

Unified File Management System
===============================
Handles all filesystem operations, file generation, and output registration.
Consolidates joi_files.py, joi_filesystem.py, and joi_file_output.py.
"""

import os
import io
import re
import json
import time
import base64
import threading
from pathlib import Path
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import joi_companion
from flask import jsonify, request as flask_req, send_file

# ── Optional Dependencies ────────────────────────────────────────────────────
try:
    from fpdf import FPDF
    HAVE_FPDF = True
except ImportError:
    HAVE_FPDF = False

try:
    from docx import Document as DocxDocument
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False

try:
    from pypdf import PdfReader
    HAVE_PYPDF = True
except ImportError:
    HAVE_PYPDF = False

# ── Lazy Imports & Helpers ───────────────────────────────────────────────────
def _require_user():
    from modules.joi_memory import require_user
    return require_user()

def _log_learning(event_type: str, data: Dict[str, Any]):
    from modules.joi_memory import log_learning_event
    log_learning_event(event_type, data)

def _download_url(filepath: Path) -> str:
    """Register the file for download and return a /download/<id> URL."""
    from modules.joi_downloads import register_download
    try:
        return register_download(filepath)
    except Exception:
        # Fallback: legacy /file/project/ URL
        try:
            rel = filepath.resolve().relative_to(BASE_DIR.resolve())
            return "/file/project/" + str(rel).replace("\\", "/")
        except ValueError:
            return "/file/project/" + filepath.name

# ── Paths & Configuration ────────────────────────────────────────────────────
BASE_DIR = Path(__file__).resolve().parent.parent

# Output directories
ASSETS_DIR = BASE_DIR / "assets"
FILES_DIR = ASSETS_DIR / "files"
OUTPUTS_DIR = ASSETS_DIR / "outputs"
PROJECTS_DIR = BASE_DIR / "projects"
RESEARCH_DIR = BASE_DIR / "research"
PROPOSALS_DIR = BASE_DIR / "proposals"

# Ensure all directories exist
for directory in [FILES_DIR, OUTPUTS_DIR, PROJECTS_DIR, RESEARCH_DIR, PROPOSALS_DIR]:
    directory.mkdir(parents=True, exist_ok=True)

# Filesystem Roots
FILE_ROOTS = {
    "project": str(BASE_DIR),
    "home": str(Path.home()),
    "desktop": str(Path.home() / "Desktop"),
    "documents": str(Path.home() / "Documents"),
    "downloads": str(Path.home() / "Downloads"),
    "pictures": str(Path.home() / "Pictures"),
    "music": str(Path.home() / "Music"),
    "videos": str(Path.home() / "Videos"),
    "appdata": str(Path.home() / "AppData"),
    "system": "C:\\",
    "programfiles": "C:\\Program Files",
    "temp": str(Path.home() / "AppData" / "Local" / "Temp"),
}

# File type definitions
TEXT_EXTS = {".txt", ".md", ".py", ".js", ".html", ".css", ".json", ".xml", ".yaml", ".yml",
             ".ini", ".cfg", ".log", ".csv", ".sh", ".bat"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".svg"}
PDF_EXTS = {".pdf"}
CODE_EXTS = {".py", ".js", ".java", ".cpp", ".c", ".cs", ".go", ".rs", ".rb", ".php"}

MAX_READ_BYTES = 5_000_000
FILE_REGISTRY_PATH = BASE_DIR / "file_registry.json"

# ── File Registry System ─────────────────────────────────────────────────────
def _load_registry() -> Dict[str, Any]:
    """Load file registry"""
    if not FILE_REGISTRY_PATH.exists():
        return {"files": [], "projects": {}, "research_topics": {}}
    try:
        return json.loads(FILE_REGISTRY_PATH.read_text(encoding='utf-8'))
    except:
        return {"files": [], "projects": {}, "research_topics": {}}

def _save_registry(registry: Dict[str, Any]):
    """Save file registry"""
    FILE_REGISTRY_PATH.write_text(json.dumps(registry, indent=2), encoding='utf-8')

def _register_file(
    filepath: Path,
    category: str,
    description: str = "",
    project_name: str = None,
    metadata: Dict[str, Any] = None
) -> str:
    """Register a file in the registry"""
    registry = _load_registry()
    file_id = f"file_{int(time.time() * 1000)}"
    
    file_record = {
        "file_id": file_id,
        "filename": filepath.name,
        "filepath": str(filepath),
        "category": category,
        "description": description,
        "project_name": project_name,
        "size_bytes": filepath.stat().st_size if filepath.exists() else 0,
        "created_at": time.time(),
        "created_datetime": datetime.now().isoformat(),
        "metadata": metadata or {}
    }
    
    registry["files"].append(file_record)
    if project_name:
        if project_name not in registry["projects"]:
            registry["projects"][project_name] = []
        registry["projects"][project_name].append(file_id)
    
    if len(registry["files"]) > 1000:
        registry["files"] = registry["files"][-1000:]
    
    _save_registry(registry)
    return file_id

# ── Helper Logic ─────────────────────────────────────────────────────────────
def resolve_path(root: str, relpath: str) -> Optional[Path]:
    if root not in FILE_ROOTS:
        return None
    root_path = Path(FILE_ROOTS[root])
    if not root_path.exists():
        return None
    try:
        # Prevent path traversal
        target = (root_path / relpath).resolve()
        # On Windows, resolve() might add drive letters or change case. 
        # But we check relative_to as a security measure.
        target.relative_to(root_path)
    except:
        return None
    return target

def _detect_language(filename: str) -> str:
    ext_to_lang = {
        ".py": "python", ".js": "javascript", ".ts": "typescript", ".jsx": "react",
        ".html": "html", ".css": "css", ".json": "json", ".md": "markdown",
        ".txt": "text", ".sh": "bash", ".sql": "sql", ".java": "java",
        ".cpp": "cpp", ".c": "c", ".go": "go", ".rs": "rust", ".php": "php", ".rb": "ruby"
    }
    ext = Path(filename).suffix.lower()
    return ext_to_lang.get(ext, "unknown")

# ── Core Operations ──────────────────────────────────────────────────────────

def fs_list(root: str, dir: str = "", pattern: str = "*", max_items: int = 100) -> Dict[str, Any]:
    """List files and directories in a root."""
    try:
        base = resolve_path(root, dir)
        if not base or not base.is_dir():
            return {"ok": False, "error": f"Invalid path: {root}/{dir}"}
        items = []
        for item in base.glob(pattern):
            if len(items) >= max_items:
                break
            try:
                rel = item.relative_to(Path(FILE_ROOTS[root]))
                items.append({
                    "name": item.name, "path": str(rel),
                    "type": "dir" if item.is_dir() else "file",
                    "size": item.stat().st_size if item.is_file() else 0,
                })
            except (PermissionError, OSError):
                continue
        return {"ok": True, "root": root, "dir": dir, "items": items, "count": len(items),
                "truncated": len(items) >= max_items}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def fs_read(root: str, path: str) -> Dict[str, Any]:
    """Read a file's content."""
    try:
        fp = resolve_path(root, path)
        if not fp or not fp.is_file():
            return {"ok": False, "error": f"Not found: {root}/{path}"}
        size = fp.stat().st_size
        if size > MAX_READ_BYTES:
            return {"ok": False, "error": f"File too large ({size} bytes)"}
        ext = fp.suffix.lower()

        if ext in TEXT_EXTS or ext in CODE_EXTS:
            return {"ok": True, "type": "text", "text": fp.read_text(encoding='utf-8', errors='replace'), "size": size}
        elif ext in PDF_EXTS:
            if not HAVE_PYPDF:
                return {"ok": False, "error": "PDF reading requires pypdf. Install with: pip install pypdf"}
            try:
                reader = PdfReader(str(fp))
                text = "\n\n".join(p.extract_text() or "" for p in reader.pages)
                return {"ok": True, "type": "pdf", "text": text, "pages": len(reader.pages), "size": size}
            except Exception as e:
                return {"ok": False, "error": f"PDF read error: {e}"}
        elif ext in IMAGE_EXTS:
            data = base64.b64encode(fp.read_bytes()).decode('utf-8')
            return {"ok": True, "type": "image", "data": f"data:image/{ext[1:]};base64,{data}", "size": size}
        return {"ok": False, "error": f"Unsupported file type: {ext}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def fs_search(root: str, dir: str = "", query: str = "", max_results: int = 30) -> Dict[str, Any]:
    """Search for files by name or content."""
    max_results = min(max_results, 50)
    try:
        base = resolve_path(root, dir)
        if not base:
            return {"ok": False, "error": f"Invalid path: {root}/{dir}"}
        results = []
        q = query.lower()
        files_scanned = 0
        for item in base.rglob("*"):
            files_scanned += 1
            if files_scanned > 5000: break # Safety cap
            if len(results) >= max_results: break
            if not item.is_file(): continue
            try:
                if q in item.name.lower():
                    results.append({"path": str(item.relative_to(Path(FILE_ROOTS[root]))),
                                    "name": item.name, "match": "filename", "size": item.stat().st_size})
                    continue
                if item.suffix.lower() in TEXT_EXTS and item.stat().st_size < 500000:
                    content = item.read_text(encoding='utf-8', errors='ignore')
                    if q in content.lower():
                        snippet = next((l[:150] for l in content.split('\n') if q in l.lower()), "")
                        results.append({"path": str(item.relative_to(Path(FILE_ROOTS[root]))),
                                        "name": item.name, "match": "content", "snippet": snippet})
            except: continue
        return {"ok": True, "root": root, "query": query, "results": results, "count": len(results)}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def generate_file(filename: str, content: str, format: str = "txt") -> Dict[str, Any]:
    """Generate a downloadable document (PDF, DOCX, TXT, MD)."""
    try:
        safe = re.sub(r'[^\w\s\-.]', '', filename).strip().replace(' ', '_')
        stem = Path(safe).stem or "output"
        
        if format == "txt":
            out = FILES_DIR / f"{stem}.txt"
            out.write_text(content, encoding='utf-8')
        elif format == "md":
            out = FILES_DIR / f"{stem}.md"
            out.write_text(content, encoding='utf-8')
        elif format == "pdf":
            if not HAVE_FPDF:
                out = FILES_DIR / f"{stem}.txt"
                out.write_text(content, encoding='utf-8')
                _log_learning("file_generated", {"format": "txt", "name": stem, "note": "fpdf fallback"})
                url = _download_url(out)
                return {"ok": True, "message": "fpdf2 not installed -- saved as TXT.", "url": url}
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=12)
            for line in content.split('\n'):
                if line.strip() == '': pdf.ln(4)
                else: pdf.multi_cell(0, 5, line)
            out = FILES_DIR / f"{stem}.pdf"
            pdf.output(str(out))
        elif format == "docx":
            if not HAVE_DOCX:
                out = FILES_DIR / f"{stem}.txt"
                out.write_text(content, encoding='utf-8')
                url = _download_url(out)
                return {"ok": True, "message": "python-docx not installed -- saved as TXT.", "url": url}
            doc = DocxDocument()
            for line in content.split('\n'):
                if line.startswith('# '): doc.add_heading(line[2:].strip(), level=1)
                elif line.startswith('## '): doc.add_heading(line[3:].strip(), level=2)
                elif line.strip() == '': doc.add_paragraph('')
                else: doc.add_paragraph(line)
            out = FILES_DIR / f"{stem}.docx"
            doc.save(str(out))
        else:
            out = FILES_DIR / f"{stem}.txt"
            out.write_text(content, encoding='utf-8')

        _log_learning("file_generated", {"format": format, "name": stem})
        url = _download_url(out)
        md_link = f"[{out.name}]({url})"
        return {"ok": True, "message": f"File ready: {md_link}", "url": url, "filename": out.name, "markdown_link": md_link}
    except Exception as e:
        return {"ok": False, "error": str(e)}

# ── Specialized Output Savers (from joi_file_output) ─────────────────────────

def save_code_file(**params) -> Dict[str, Any]:
    """Save code to a file and register it."""
    _require_user()
    code = params.get("code", "")
    filename = params.get("filename", "code.py")
    description = params.get("description", "")
    project_name = params.get("project_name")
    destination = params.get("destination", "outputs")
    
    if not code: return {"ok": False, "error": "No code provided"}
    
    dest_dir = OUTPUTS_DIR
    if destination == "projects":
        dest_dir = PROJECTS_DIR
        if project_name:
            dest_dir = dest_dir / project_name
            dest_dir.mkdir(parents=True, exist_ok=True)
    elif destination == "proposals":
        dest_dir = PROPOSALS_DIR
    
    filepath = dest_dir / filename
    try:
        filepath.write_text(code, encoding='utf-8')
        file_id = _register_file(filepath, category="code", description=description, 
                                project_name=project_name, 
                                metadata={"language": _detect_language(filename), "lines": code.count('\n')+1})
        url = _download_url(filepath)
        return {"ok": True, "file_id": file_id, "url": url, "message": f"Code saved as {filename}. Download: {url}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def save_text_file(**params) -> Dict[str, Any]:
    """Save text content to a file and register it."""
    _require_user()
    content = params.get("content", "")
    filename = params.get("filename", "document.txt")
    description = params.get("description", "")
    fmt = params.get("format", "txt")
    destination = params.get("destination", "outputs")
    
    if not content: return {"ok": False, "error": "No content provided"}
    if not filename.endswith(f".{fmt}"): filename = f"{filename}.{fmt}"
    
    dest_dir = OUTPUTS_DIR if destination == "outputs" else RESEARCH_DIR
    filepath = dest_dir / filename
    
    try:
        filepath.write_text(content, encoding='utf-8')
        file_id = _register_file(filepath, category="document", description=description,
                                metadata={"format": fmt, "words": len(content.split())})
        url = _download_url(filepath)
        return {"ok": True, "file_id": file_id, "url": url, "message": f"Document saved. Download: {url}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

def save_research_findings(**params) -> Dict[str, Any]:
    """Format and save research findings."""
    _require_user()
    topic = params.get("topic", "Research")
    findings = params.get("findings", "")
    sources = params.get("sources", [])
    summary = params.get("summary", "")
    
    if not findings: return {"ok": False, "error": "No findings provided"}
    
    safe_topic = re.sub(r'[^\w\s-]', '', topic).strip().replace(' ', '_')
    filename = f"research_{safe_topic}_{int(time.time())}.md"
    
    content = f"# Research: {topic}\n**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
    if summary: content += f"## Summary\n{summary}\n\n"
    content += f"## Findings\n{findings}\n\n"
    if sources:
        content += "## Sources\n"
        for i, s in enumerate(sources, 1): content += f"{i}. {s}\n"
        
    filepath = RESEARCH_DIR / filename
    try:
        filepath.write_text(content, encoding='utf-8')
        registry = _load_registry()
        if topic not in registry["research_topics"]: registry["research_topics"][topic] = []
        file_id = _register_file(filepath, category="research", description=f"Research: {topic}")
        registry["research_topics"][topic].append(file_id)
        _save_registry(registry)
        url = _download_url(filepath)
        return {"ok": True, "file_id": file_id, "url": url, "topic": topic, "message": f"Research saved. Download: {url}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}

# ── Registry Retrieval ───────────────────────────────────────────────────────

def list_registry_files(**params) -> Dict[str, Any]:
    """List files registered in the system."""
    _require_user()
    cat = params.get("category")
    proj = params.get("project_name")
    limit = params.get("limit", 50)
    
    registry = _load_registry()
    files = registry["files"]
    if cat: files = [f for f in files if f["category"] == cat]
    if proj: files = [f for f in files if f.get("project_name") == proj]
    
    files.sort(key=lambda x: x["created_at"], reverse=True)
    return {"ok": True, "files": files[:limit], "total": len(files)}

def get_registered_content(**params) -> Dict[str, Any]:
    """Read content of a registered file by ID."""
    _require_user()
    fid = params.get("file_id")
    if not fid: return {"ok": False, "error": "file_id required"}
    
    registry = _load_registry()
    record = next((f for f in registry["files"] if f["file_id"] == fid), None)
    if not record: return {"ok": False, "error": "Not registered"}
    
    path = Path(record["filepath"])
    if not path.exists(): return {"ok": False, "error": "File gone"}
    try:
        return {"ok": True, "filename": record["filename"], "content": path.read_text(encoding='utf-8'), "metadata": record}
    except Exception as e: return {"ok": False, "error": str(e)}

# ── Tool Registration ────────────────────────────────────────────────────────

joi_companion.register_tool(
    {"type": "function", "function": {
        "name": "generate_file",
        "description": "Create downloadable document (PDF, TXT, DOCX, MD). Returns a markdown link.",
        "parameters": {"type": "object", "properties": {
            "filename": {"type": "string"},
            "content": {"type": "string"},
            "format": {"type": "string", "enum": ["txt", "md", "pdf", "docx"], "default": "txt"}
        }, "required": ["filename", "content"]}
    }},
    generate_file
)

joi_companion.register_tool(
    {"type": "function", "function": {
        "name": "fs_list",
        "description": "List files in: project, home, desktop, documents, downloads, etc.",
        "parameters": {"type": "object", "properties": {
            "root": {"type": "string", "enum": list(FILE_ROOTS.keys())},
            "dir": {"type": "string", "default": ""},
            "pattern": {"type": "string", "default": "*"}
        }, "required": ["root"]}
    }},
    fs_list
)

joi_companion.register_tool(
    {"type": "function", "function": {
        "name": "fs_read",
        "description": "Read file (txt, py, js, pdf, img).",
        "parameters": {"type": "object", "properties": {
            "root": {"type": "string", "enum": list(FILE_ROOTS.keys())},
            "path": {"type": "string"}
        }, "required": ["root", "path"]}
    }},
    fs_read
)

joi_companion.register_tool(
    {"type": "function", "function": {
        "name": "save_code_file",
        "description": "Save complete code files.",
        "parameters": {"type": "object", "properties": {
            "code": {"type": "string"},
            "filename": {"type": "string"},
            "description": {"type": "string"},
            "project_name": {"type": "string"},
            "destination": {"type": "string", "enum": ["outputs", "projects", "proposals"]}
        }, "required": ["code", "filename"]}
    }},
    save_code_file
)

joi_companion.register_tool(
    {"type": "function", "function": {
        "name": "save_research_findings",
        "description": "Save research with sources.",
        "parameters": {"type": "object", "properties": {
            "topic": {"type": "string"},
            "findings": {"type": "string"},
            "sources": {"type": "array", "items": {"type": "string"}},
            "summary": {"type": "string"}
        }, "required": ["topic", "findings"]}
    }},
    save_research_findings
)

# ── Flask Routes ─────────────────────────────────────────────────────────────

def _fs_browse_route():
    _require_user()
    return jsonify(fs_list(flask_req.args.get("root", "project"), flask_req.args.get("dir", "")))

def _fs_read_route():
    _require_user()
    data = flask_req.get_json(silent=True) or {}
    return jsonify(fs_read(data.get("root", "project"), data.get("path", "")))

def _files_api_route():
    _require_user()
    if flask_req.method == "GET":
        return jsonify(list_registry_files(category=flask_req.args.get("category"), 
                                          project_name=flask_req.args.get("project")))
    data = flask_req.get_json(silent=True) or {}
    action = data.get("action")
    if action == "save_code": return jsonify(save_code_file(**data))
    if action == "save_text": return jsonify(save_text_file(**data))
    if action == "get_content": return jsonify(get_registered_content(**data))
    return jsonify({"ok": False, "error": "Invalid action"})

joi_companion.register_route("/fs/browse", ["GET"], _fs_browse_route, "fs_browse")
joi_companion.register_route("/fs/read", ["POST"], _fs_read_route, "fs_read_post")
joi_companion.register_route("/files", ["GET", "POST"], _files_api_route, "files_api")
